from dataclasses import dataclass, field
from pathlib import Path


SUPPORTED_EXTENSIONS = {".pdf", ".md", ".txt", ".docx", ".epub"}


@dataclass
class Document:
    content: str
    source: str
    metadata: dict = field(default_factory=dict)


def parse_file(path: Path) -> Document:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _parse_pdf(path)
    elif suffix in (".md", ".txt"):
        return _parse_text(path)
    elif suffix == ".docx":
        return _parse_docx(path)
    elif suffix == ".epub":
        return _parse_epub(path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


def _parse_text(path: Path) -> Document:
    content = path.read_text(encoding="utf-8", errors="replace")
    return Document(
        content=content,
        source=str(path),
        metadata={"type": path.suffix.lstrip(".")},
    )


def _parse_pdf(path: Path) -> Document:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    content = "\n\n".join(p for p in pages if p.strip())
    return Document(
        content=content,
        source=str(path),
        metadata={"type": "pdf", "pages": len(pages)},
    )


def _parse_docx(path: Path) -> Document:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    content = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return Document(content=content, source=str(path), metadata={"type": "docx"})


def _parse_epub(path: Path) -> Document:
    import ebooklib
    from ebooklib import epub
    from bs4 import BeautifulSoup

    book = epub.read_epub(str(path))
    chapters = []
    for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
        soup = BeautifulSoup(item.get_content(), "html.parser")
        text = soup.get_text()
        if text.strip():
            chapters.append(text)
    content = "\n\n".join(chapters)
    return Document(content=content, source=str(path), metadata={"type": "epub"})


def collect_files(path: Path) -> list[Path]:
    if path.is_file():
        return [path] if path.suffix.lower() in SUPPORTED_EXTENSIONS else []
    return sorted(f for f in path.rglob("*") if f.suffix.lower() in SUPPORTED_EXTENSIONS)
